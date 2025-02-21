from typing import List
import asyncio
from aiogram import types
from aiogram.dispatcher.filters import Text
from data.config import ADMINS
from keyboards.inline.admin_keys import (admin_panel_keyboard, admin_back_keyboard, channels_list_builder, 
                                        channel_details_keyboard, back_channels_list,
                                        admin_books_button, admin_lessons_button, admin_subjects_button,
                                        admin_back_lessons_books, admin_back_lessons_books_2, admin_back_book_list, advert_type_keyboard,
                                        get_tests_button, show_tests_pages, shuffling_type_keyboard, test_details_keyboard)
from loader import dp, channel, bot, booksdb, lessonsdb, subjectsdb, db, testsdb
from filters.private_chat_filter import IsPrivate
from aiogram.dispatcher import FSMContext
from aiogram.utils.exceptions import ChatNotFound, BadRequest
from aiogram_media_group import media_group_handler
from aiogram.types import InputMediaDocument, InputMediaPhoto, InputMediaVideo, InputMediaAudio
from aiogram import types
from aiogram.dispatcher import FSMContext
from docx import Document
from io import BytesIO
from loader import db
from aiogram.types import ReplyKeyboardRemove
from datetime import datetime
import pytz
@dp.message_handler(IsPrivate(), content_types=types.ContentType.ANY, state="*", commands="admin", user_id=ADMINS)
async def admin_panel(message: types.Message, state: FSMContext):
    await state.finish()
    await message.answer("🔑 Assalomu alaykum admin! Admin panelga xush kelibsiz! Pastga tugmalar orqali menuni tanlang!", reply_markup=admin_panel_keyboard)


@dp.callback_query_handler(IsPrivate(), text="admin:back", state="*", user_id=ADMINS)
async def admin_back(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("🔑 Assalomu alaykum admin! Admin panelga xush kelibsiz! Pastga tugmalar orqali menuni tanlang!", reply_markup=admin_panel_keyboard)


@dp.callback_query_handler(IsPrivate(), text="admin:subscription", user_id=ADMINS)
async def admin_subscription(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    channels_list = channel.get_channels()
    await callback.message.edit_text("🔐 Majburiy obunaga qo'yilgan kanallar ro'yxati", reply_markup=await channels_list_builder(channels_list))

@dp.callback_query_handler(IsPrivate(), text="admin:back_channels_list", state="*", user_id=ADMINS)
async def admin_back_channels_list(callback: types.CallbackQuery, state: FSMContext):
    await admin_subscription(callback, state)

@dp.callback_query_handler(IsPrivate(), text_contains="admin:channel", state="*", user_id=ADMINS)
async def admin_channel(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    username = callback.data.split(":")[2]
    chat = await bot.get_chat(username)
    time = channel.get_time_channel(username)
    members_count = await bot.get_chat_members_count(username)

    await callback.message.edit_text(f"<b>🔐 Kanal:</b> <a href='tg://user?id={chat.id}'>{chat.title}</a>\n\n<b>🔗 Ulangan:</b> {time}\n\n<b>👥 Obunachilar:</b> {members_count}", reply_markup=channel_details_keyboard(username))

@dp.callback_query_handler(IsPrivate(), text="admin:add_channel", state="*", user_id=ADMINS)
async def admin_add_channel(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.edit_text("🔐 Kanal usernameini kiriting:", reply_markup=back_channels_list)
    await state.set_state("admin:add_channel")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.ANY, state="admin:add_channel", user_id=ADMINS)
async def admin_add_channel_message(message: types.Message, state: FSMContext):
    username = message.text
    username = username.replace("@", "")
    username = username.replace("https://t.me/", "")
    username = "@" + username
    username = username.lower()
    try:
        bot_member = await bot.get_chat_member(username, bot.id)
        channel.save_channel(username)
        channels_list = channel.get_channels()
        await message.answer("<b>🔐 Kanal ro'yxatga muvaffaqiyatli qo'shildi!</b>", reply_markup=await channels_list_builder(channels_list))
        await state.finish()
    except ChatNotFound:
        await message.answer("<b>🔐 Kanal usernameini noto'g'ri kiritdingiz yoki bot kanalda admin emas!</b>\n\n♻️ Qaytadan kiritishga harakat qilib ko'ring:", reply_markup=back_channels_list)
    except BadRequest:
        await message.answer("<b>🤖 Botga foydalanuvchilar boshqarish ruxsati berilmagan!</b>\n\n♻️ Qaytadan kiritishga harakat qilib ko'ring:", reply_markup=back_channels_list)
    except:
        await message.answer("<b>✅ Bu kanal allaqachon ro'yxatga qo'shilgan!</b>\n\n♻️ Qaytadan kiritishga harakat qilib ko'ring:", reply_markup=back_channels_list)

@dp.callback_query_handler(IsPrivate(), text_contains="admin:delete_channel", state="*", user_id=ADMINS)
async def admin_delete_channel(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    username = callback.data.split(":")[2]  
    channel.del_channel(username)
    channels_list = channel.get_channels()
    await callback.message.edit_text("🔐 Kanal ro'yxatdan olib tashlandi!", reply_markup=await channels_list_builder(channels_list))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:action", state="*", user_id=ADMINS)
async def admin_action(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    action = callback.data.split(":")[2]
    subject = callback.data.split(":")[3]
    if action == "tests":
        await admin_tests_type_action(callback.message, subject)
    elif action == "books":
        await admin_books_action(callback.message, subject)
    elif action == "lessons":
        await admin_lessons_action(callback.message, subject)



# Books
async def admin_books_action(message: types.Message, subject: str, action_type = "edit"):
    if action_type == "edit":
        await message.edit_text(f"📕 <b>{subject}</b>, kitobni tanlang:", reply_markup=admin_books_button(subject, booksdb.get_books(subject)))
    elif action_type == "send":
        await message.answer(f"📕 <b>{subject}</b>, kitobni tanlang:", reply_markup=admin_books_button(subject, booksdb.get_books(subject)))

# Lessons
async def admin_lessons_action(message: types.Message, subject: str, action_type = "edit"):
    if action_type == "edit":
        await message.edit_text(f"📕 <b>{subject}</b>, kitobni tanlang:", reply_markup=admin_lessons_button(subject, lessonsdb.get_lessons(subject)))
    elif action_type == "send":
        await message.answer(f"📕 <b>{subject}</b>, kitobni tanlang:", reply_markup=admin_lessons_button(subject, lessonsdb.get_lessons(subject)))

# Tests
async def admin_tests_type_action(message: types.Message, subject: str, action_type = "edit"):
    if action_type == "edit":
        await message.edit_text(f"📚 Testlar ro'yxatidan kerakli testni tanlang yoki yangi test yarating!", reply_markup=get_tests_button(subject, testsdb.get_tests(subject=subject, limit=5, page=1, test_type="all")))
    elif action_type == "send":
        await message.answer(f"📚 Testlar ro'yxatidan kerakli testni tanlang yoki yangi test yarating!", reply_markup=get_tests_button(subject, testsdb.get_tests(subject=subject, limit=5, page=1, test_type="all")))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:add_subject", state="*", user_id=ADMINS)
async def admin_add_subject(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("📚 Fan nomini kiriting(Eslatma bu fan barcha turdagi fanlar bo'limiga qo'shiladi):", reply_markup=admin_back_keyboard)
    await state.set_state("admin:add_subject")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:add_subject", user_id=ADMINS)
async def admin_add_subject_message(message: types.Message, state: FSMContext):
    subject = message.text
    subjectsdb.save_subject(subject)
    await message.answer("📚 Fan ro'yxatga muvaffaqiyatli qo'shildi!", reply_markup=admin_back_keyboard)
    await state.finish()





# books section
async def admin_books(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("📚 Darsliklar sozlamalarini o'zgartirish uchun fanni tanlang!", reply_markup=admin_subjects_button(action="books", subjects=subjectsdb.get_subjects()))

async def admin_lessons(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("📚 Darsliklar sozlamalarini o'zgartirish uchun fanni tanlang!", reply_markup=admin_subjects_button(action="lessons", subjects=subjectsdb.get_subjects()))


@dp.callback_query_handler(IsPrivate(), text_contains="admin:back_subjects", state="*", user_id=ADMINS)
async def admin_back_subjects(callback: types.CallbackQuery, state: FSMContext):
    action = callback.data.split(":")[2]
    if action == "books":
        await admin_books(callback, state)
    elif action == "lessons":
        await admin_lessons(callback, state)
    elif action == "tests":
        await admin_tests(callback, state)

@dp.callback_query_handler(IsPrivate(), text_contains="admin:delete_subject", state="*", user_id=ADMINS)
async def admin_delete_subject(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    subjectsdb.del_subject(subject)
    await callback.message.edit_text("📚 Fan ro'yxatdan olib tashlandi!", reply_markup=admin_subjects_button(action="books", subjects=subjectsdb.get_subjects()))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:back_book_lessons", state="*", user_id=ADMINS)
async def admin_back_book_lessons(callback: types.CallbackQuery, state: FSMContext):
    action = callback.data.split(":")[2]
    subject = callback.data.split(":")[3]
    await state.finish()
    await callback.message.delete()
    if action == "books":
        await admin_books_action(callback.message, subject, "send")
    elif action == "lessons":
        await admin_lessons_action(callback.message, subject, "send")

@dp.callback_query_handler(IsPrivate(), text_contains="admin:books", state="*", user_id=ADMINS)
async def admin_books(callback: types.CallbackQuery, state: FSMContext):
    try:await state.finish()
    except:pass

    await callback.message.edit_text("📕 Darsliklar sozlamalarini o'zgartirish uchun fanni tanlang!", reply_markup=admin_subjects_button(action="books", subjects=subjectsdb.get_subjects()))


@dp.callback_query_handler(IsPrivate(), text_contains="admin:get_book", state="*", user_id=ADMINS)
async def admin_get_book(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    book = callback.data.split(":")[3]
    file_id = booksdb.get_book(book, subject)[0]
    message_id = booksdb.get_book(book, subject)[1]
    await callback.message.delete()
    try:
        await callback.message.answer_document(file_id, caption=f"📕 <b>{subject}</b>dan <b>{book}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "books", book))
    except:
        await bot.copy_message(message_id=message_id, chat_id=callback.from_user.id, from_chat_id="-1002256408723", caption=f"📕 <b>{subject}</b>dan <b>{book}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "books", book))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:edit_book", state="*", user_id=ADMINS)
async def admin_edit_book(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.delete()
    await state.finish()
    subject = callback.data.split(":")[2]
    book = callback.data.split(":")[3]
    await callback.message.answer("📕 <b>Menga yangi kitobni yuboring!</b>", reply_markup=admin_back_lessons_books_2(subject, "books"))
    await state.set_data({"subject": subject, "book": book})
    await state.set_state("admin:edit_book")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:edit_book", user_id=ADMINS)
async def admin_edit_book_message(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    book = data["book"]
    await bot.copy_message(message_id=message.message_id, chat_id=message.chat.id, from_chat_id=message.from_user.id, caption=f"📕 <b>{book}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "books", book))
    document = await bot.copy_message(message_id=message.message_id, chat_id="-1002256408723", from_chat_id=message.from_user.id, caption=f"Kitob BAZA!")
    message_id = document.message_id
    booksdb.edit_book(book, subject, message.document.file_id, message_id)

@dp.callback_query_handler(IsPrivate(), text_contains="admin:add_book", state="*", user_id=ADMINS)
async def admin_add_book(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    await callback.message.edit_text("📕 Kitob nomini kiriting:", reply_markup=admin_back_book_list(subject, "books"))
    await state.set_data({"subject": subject})
    await state.set_state("admin:add_book")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:add_book", user_id=ADMINS)
async def admin_add_book_message(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    book = message.text
    await message.answer(f"📕 <b>{book}</b>, kitobi mediani yuboring", reply_markup=admin_back_book_list(subject, "books"))
    await state.set_data({"subject": subject, "book": book})
    await state.set_state("admin:add_book_document")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:add_book_document", user_id=ADMINS)
async def admin_add_book_document(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    book = data["book"]
    await message.answer(f"📕 <b>{book}</b>, kitobi media yuklanmoqda...")
    await bot.copy_message(message_id=message.message_id, chat_id=message.chat.id, from_chat_id=message.from_user.id, caption=f"📕 <b>{book}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "books", book))
    document = await bot.copy_message(message_id=message.message_id, chat_id="-1002256408723", from_chat_id=message.from_user.id, caption=f"Kitob BAZA!")
    message_id = document.message_id
    booksdb.save_book(book, subject, message.document.file_id, message_id)


# lessons section
@dp.callback_query_handler(IsPrivate(), text="admin:lessons", state="*", user_id=ADMINS)
async def admin_lessons(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("📒 Dars ishlanmalari", reply_markup=admin_subjects_button(action="lessons", subjects=subjectsdb.get_subjects()))


@dp.callback_query_handler(IsPrivate(), text_contains="admin:get_lesson", state="*", user_id=ADMINS)
async def admin_get_lesson(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    lesson = callback.data.split(":")[3]
    file_id = lessonsdb.get_lesson(lesson, subject)[0]
    message_id = lessonsdb.get_lesson(lesson, subject)[1]
    await callback.message.delete()
    try:
        await callback.message.answer_document(file_id, caption=f"📕 <b>{subject}</b>dan <b>{lesson}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "lessons", lesson))
    except:
        await bot.copy_message(message_id=message_id, chat_id=callback.from_user.id, from_chat_id="-1002256408723", caption=f"📕 <b>{subject}</b>dan <b>{lesson}</b>, kitobi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "lessons", lesson))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:edit_lesson", state="*", user_id=ADMINS)
async def admin_edit_lesson(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.delete()
    await state.finish()
    subject = callback.data.split(":")[2]
    lesson = callback.data.split(":")[3]
    await callback.message.answer("📕 <b>Menga yangi dars ishlanmasini yuboring!</b>", reply_markup=admin_back_lessons_books_2(subject, "lessons"))
    await state.set_data({"subject": subject, "lesson": lesson})
    await state.set_state("admin:edit_lesson")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:edit_lesson", user_id=ADMINS)
async def admin_edit_lesson_message(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    lesson = data["lesson"]
    await bot.copy_message(message_id=message.message_id, chat_id=message.chat.id, from_chat_id=message.from_user.id, caption=f"📕 <b>{lesson}</b>, darsi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "lessons", lesson))
    document = await bot.copy_message(message_id=message.message_id, chat_id="-1002256408723", from_chat_id=message.from_user.id, caption=f"Kitob BAZA!")
    message_id = document.message_id
    lessonsdb.edit_lesson(lesson, subject, message.document.file_id, message_id)



@dp.callback_query_handler(IsPrivate(), text_contains="admin:add_lesson", state="*", user_id=ADMINS)
async def admin_add_lesson(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    await callback.message.edit_text("📕 Kitob nomini kiriting:", reply_markup=admin_back_book_list(subject, "lessons"))
    await state.set_data({"subject": subject})
    await state.set_state("admin:add_lesson")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:add_lesson", user_id=ADMINS)
async def admin_add_lesson_message(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    lesson = message.text
    await message.answer(f"📕 <b>{lesson}</b>, darsi mediani yuboring", reply_markup=admin_back_book_list(subject, "lessons"))
    await state.set_data({"subject": subject, "lesson": lesson})
    await state.set_state("admin:add_lesson_document")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:add_lesson_document", user_id=ADMINS)
async def admin_add_lesson_document(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    lesson = data["lesson"]
    await message.answer(f"📕 <b>{lesson}</b>, darsi media yuklanmoqda...")
    await bot.copy_message(message_id=message.message_id, chat_id=message.chat.id, from_chat_id=message.from_user.id, caption=f"📕 <b>{lesson}</b>, darsi mediani pastdagi tugma orqali alishtirishingiz mumkin !", reply_markup=admin_back_lessons_books(subject, "lessons", lesson))
    document = await bot.copy_message(message_id=message.message_id, chat_id="-1002256408723", from_chat_id=message.from_user.id, caption=f"Kitob BAZA!")
    message_id = document.message_id
    lessonsdb.save_lesson(lesson, subject, message.document.file_id, message_id)  

@dp.callback_query_handler(IsPrivate(), text="admin:announcement", state="*", user_id=ADMINS)
async def admin_announcement(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("💬 Reklama tarqatish uchun uning turini ko'rsating:", reply_markup=advert_type_keyboard)

@dp.callback_query_handler(IsPrivate(), text_contains="admin:advert_type", state="*", user_id=ADMINS)
async def admin_advert_type(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.delete()
    await state.finish()
    advert_type = callback.data.split(":")[2]
    if advert_type == "forward":
        await forward_advert(callback.message, state)
    elif advert_type == "simple":
        await simple_advert(callback.message, state)

async def forward_advert(message: types.Message, state: FSMContext):
    await message.answer("💬 Forward qilinishi kerak bo'lgan postni menga yuboring:", reply_markup=admin_back_keyboard)
    await state.set_state("admin:advert_forward")

async def simple_advert(message: types.Message, state: FSMContext):
    await message.answer("💬 Reklamani menga yuboring (Reklama foto, video, audio, dokument yoki tekst bo'lishi mumkin):", reply_markup=admin_back_keyboard)
    await state.set_state("admin:advert_simple")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.ANY, state="admin:advert_forward", user_id=ADMINS)
async def admin_advert_forward(message: types.Message, state: FSMContext):
    await state.finish()
    users = db.select_all_users()
    x = 0
    y = 0
    start = datetime.now(pytz.timezone('Asia/Tashkent'))
    i = await message.answer("🔰 Reklama yuborilmoqda, iltimos biroz kutib turing...")
    await message.answer("📢 Reklama tarqatilmoqda...")
    for user in users:
        try:
            await bot.forward_message(chat_id=user[0], from_chat_id=message.from_user.id, message_id=message.message_id)
            y+=1
        except:
            pass
            x+=1


        await asyncio.sleep(0.1)

    finish = datetime.now(pytz.timezone('Asia/Tashkent'))
    farq = finish - start
    
    await message.answer(f"<b>📣 Reklama yuborildi</b>\n\n"
                         f"✅ Qabul qildi: {y} ta\n"
                         f"❌ Yuborilmadi: {x} ta\n\n"
                         f"<b>⏰ Boshlandi:</b> {start.strftime('%H:%M:%S')}\n"
                         f"<b>⏰ Yakunlandi:</b> {finish.strftime('%H:%M:%S')}\n\n"
                         f"<b>🕓 Umumiy ketgan vaqt:</b> {farq.seconds} soniya", reply_markup=admin_back_keyboard)

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:advert_simple", user_id=ADMINS)
async def admin_advert_simple(message: types.Message, state: FSMContext):
    await state.finish()
    users = db.select_all_users()
    x = 0
    y = 0
    start = datetime.now(pytz.timezone('Asia/Tashkent'))
    i = await message.answer("🔰 Reklama yuborilmoqda, iltimos biroz kutib turing...")
    await message.answer("📢 Reklama tarqatilmoqda...")
    for user in users:
        try:
            await bot.send_message(chat_id=user[0], text=message.text)
            y+=1
        except:
            x+=1
        await asyncio.sleep(0.1)
    finish = datetime.now(pytz.timezone('Asia/Tashkent'))
    farq = finish - start
    
    await message.answer(f"<b>📣 Reklama yuborildi</b>\n\n"
                         f"✅ Qabul qildi: {y} ta\n"
                         f"❌ Yuborilmadi: {x} ta\n\n"
                         f"<b>⏰ Boshlandi:</b> {start.strftime('%H:%M:%S')}\n"
                         f"<b>⏰ Yakunlandi:</b> {finish.strftime('%H:%M:%S')}\n\n"
                         f"<b>🕓 Umumiy ketgan vaqt:</b> {farq.seconds} soniya", reply_markup=admin_back_keyboard)

@dp.message_handler(IsPrivate(), content_types=types.ContentType.ANY, state="admin:advert_simple", user_id=ADMINS)
async def admin_advert_simple(message: types.Message, state: FSMContext):



    users = db.select_all_users()
    x = 0
    y = 0
    start = datetime.now(pytz.timezone('Asia/Tashkent'))
    i = await message.answer("🔰 Reklama yuborilmoqda, iltimos biroz kutib turing...")
    await message.answer("📢 Reklama tarqatilmoqda...")
    for user in users:
        try:
            if message.document:
                await bot.send_document(chat_id=user[0], document=message.document.file_id, caption=message.caption or '')
                
            elif message.photo:
                await bot.send_photo(chat_id=user[0], photo=message.photo[-1].file_id, caption=message.caption or '' )
                
                    
            elif message.video:
                await bot.send_video(chat_id=user[0], video=message.video.file_id, caption=message.caption or '' )

            elif message.audio:
                await bot.send_audio(chat_id=user[0], audio=message.audio.file_id, caption=message.caption or '' )
                
            y+=1
        except Exception as e: 
            x+=1
        await asyncio.sleep(0.1)
    finish = datetime.now(pytz.timezone('Asia/Tashkent'))
    farq = finish - start
    
    await message.answer(f"<b>📣 Reklama yuborildi</b>\n\n"
                         f"✅ Qabul qildi: {y} ta\n"
                         f"❌ Yuborilmadi: {x} ta\n\n"
                         f"<b>⏰ Boshlandi:</b> {start.strftime('%H:%M:%S')}\n"
                         f"<b>⏰ Yakunlandi:</b> {finish.strftime('%H:%M:%S')}\n\n"
                         f"<b>🕓 Umumiy ketgan vaqt:</b> {farq.seconds} soniya", reply_markup=admin_back_keyboard)


@dp.message_handler(IsPrivate(), content_types=types.ContentType.ANY, state="admin:advert_simple", user_id=ADMINS)
@media_group_handler
async def admin_advert_simple(messages: List[types.Message], state: FSMContext):
    await state.finish()
    media_group = []

    if isinstance(messages, list):
        pass
    else:
        messages = [messages]
    for message in messages:
        if message.document:
            media_group.append(
                InputMediaDocument(
                    media=message.document.file_id,
                    caption=message.caption or '' 
                )
            )
        elif message.photo:
            media_group.append(
                InputMediaPhoto(
                    media=message.photo[-1].file_id,
                    caption=message.caption or '' 
                )
            )
        elif message.video:
            media_group.append(
                InputMediaVideo(
                    media=message.video.file_id,
                    caption=message.caption or '' 
                )
            )
        elif message.audio:
            media_group.append(
                InputMediaAudio(
                    media=message.audio.file_id,
                    caption=message.caption or '' 
                )
            )

    users = db.select_all_users()
    x = 0
    y = 0
    start = datetime.now(pytz.timezone('Asia/Tashkent'))
    i = await message.answer("🔰 Reklama yuborilmoqda, iltimos biroz kutib turing...")
    await message.answer("📢 Reklama tarqatilmoqda...")
    for user in users:
        try:
            await bot.send_media_group(chat_id=user[0], media=media_group)
            y+=1
        except:
            x+=1
        await asyncio.sleep(0.1)
    finish = datetime.now(pytz.timezone('Asia/Tashkent'))
    farq = finish - start
    
    await message.answer(f"<b>📣 Reklama yuborildi</b>\n\n"
                         f"✅ Qabul qildi: {y} ta\n"
                         f"❌ Yuborilmadi: {x} ta\n\n"
                         f"<b>⏰ Boshlandi:</b> {start.strftime('%H:%M:%S')}\n"
                         f"<b>⏰ Yakunlandi:</b> {finish.strftime('%H:%M:%S')}\n\n"
                         f"<b>🕓 Umumiy ketgan vaqt:</b> {farq.seconds} soniya", reply_markup=admin_back_keyboard)

        

@dp.callback_query_handler(IsPrivate(), text="admin:tests", state="*", user_id=ADMINS)
async def admin_tests(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await callback.message.edit_text("📚 Test yaratishdan avval fanni tanlang !", reply_markup=admin_subjects_button(action="tests", subjects=subjectsdb.get_subjects()))

@dp.callback_query_handler(IsPrivate(), text_contains="action:tests", state="*", user_id=ADMINS)
async def admin_tests_action(callback: types.CallbackQuery, state: FSMContext):
    subject = callback.data.split(":")[2]
    await callback.message.edit_text("📚 Testlar ro'yxatidan kerakli testni tanlang yoki yangi test yarating!", reply_markup=get_tests_button(subject, testsdb.get_tests(subject=subject, limit=5, page=1, test_type="all")))
    await state.finish()


@dp.callback_query_handler(IsPrivate(), text_contains="admin:paginate_shmldrq", state="*", user_id=ADMINS)
async def admin_paginate_tests(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    page = callback.data.split(":")[3]
    if page == 'no_page':
        await callback.answer("❌ Bu sahifaga o'tib bo'lmaydi! Limitga yetildi!")
    else:
        await callback.message.edit_text("📚 Testlar ro'yxatidan kerakli testni tanlang yoki yangi test yarating!", reply_markup=get_tests_button(subject, testsdb.get_tests(subject=subject, limit=5, page=page, test_type="all")))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:shmldrq", state="*", user_id=ADMINS)
async def admin_pages(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    await callback.message.edit_text("Sahifani tanlang!", reply_markup=show_tests_pages(subject, testsdb.get_tests(subject=subject, test_type="all")))


@dp.callback_query_handler(IsPrivate(), text_contains="admin:get_shmldrq", state="*", user_id=ADMINS)
async def admin_get_test(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    test = callback.data.split(":")[3]
    test_details = testsdb.get_information(test)
    if test_details['type'] == "themed":
        type = "Mavzulashtirilgan"
    else:
        type = "Attestatsiya"
    await callback.message.edit_text(f"<b>✍️ Test nomi:</b> {test_details['test_name']}\n<b>🔀 Test aralashtirish turi:</b> {test_details['shuffling_type']}\n<b>📍 Test savollar soni:</b> {test_details['total_questions']}\n<b>📅 Test yaratilgan sana va vaqt:</b> {test_details['created_time']}\n<b>⏰ Test vaqti:</b> {test_details['timer']}\n📝 Test turi: {type}\n\n<b>🔗 Testni tarqatish uchun link:</b> https://t.me/PedagogUzbot?start=test{test}", reply_markup=test_details_keyboard(subject, test))




@dp.callback_query_handler(IsPrivate(), text_contains="admin:back_tests_page", state="*", user_id=ADMINS)
async def admin_back_questions_page(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    await admin_tests_action(callback, state)



@dp.callback_query_handler(IsPrivate(), text_contains="admin:create_test", state="*", user_id=ADMINS)
async def admin_create_test(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    await callback.message.edit_text("Test nomini kiriting:", reply_markup=admin_back_book_list(subject, "tests"))
    await state.set_data({"subject": subject})
    await state.set_state("admin:create_test")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:create_test", user_id=ADMINS)
async def admin_create_test_message(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    test_name = message.text
    await message.answer("Testni aralashtirish turini tanlang:", reply_markup=shuffling_type_keyboard)
    await state.set_data({"subject": subject, "test_name": test_name})
    await state.set_state("admin:create_test_shuffling_type")

@dp.message_handler(IsPrivate(), state="admin:create_test_shuffling_type", user_id=ADMINS)
async def admin_create_test_shuffling_type(message: types.Message, state: FSMContext):
    shuffling_type = message.text
    if shuffling_type not in ["🔀 Hammasini aralashtirish", "🔀 Savolni aralashtirish", "🔀 Javobni aralashtirish", "🔀 Umuman aralashtirmaslik"]:
        await message.answer("Testni aralashtirish turini tanlang!", reply_markup=shuffling_type_keyboard)
        return
    elif shuffling_type == "🔀 Hammasini aralashtirish":
        shuffling_type = "shuffle_all"
    elif shuffling_type == "🔀 Savolni aralashtirish":
        shuffling_type = "shuffle_question"
    elif shuffling_type == "🔀 Javobni aralashtirish":
        shuffling_type = "shuffle_answer"
    elif shuffling_type == "🔀 Umuman aralashtirmaslik":
        shuffling_type = "no_shuffle"
    data = await state.get_data()
    subject = data["subject"]
    test_name = data["test_name"]
    await state.set_data({"subject": subject, "test_name": test_name, "shuffling_type": shuffling_type})
    await state.set_state("admin:create_test_time")
    await message.answer("Testlar orasidagi vaqtni kiriting (sekundda):", reply_markup=ReplyKeyboardRemove())

@dp.message_handler(IsPrivate(), state="admin:create_test_time", user_id=ADMINS)
async def admin_create_test_shuffling_type(message: types.Message, state: FSMContext):
    data = await state.get_data()
    timer = message.text
    shuffling_type = data['shuffling_type']
    subject = data["subject"]
    test_name = data["test_name"]
    await state.set_data({"subject": subject, "test_name": test_name, "shuffling_type": shuffling_type, 'timer': timer})
    keyboard = types.ReplyKeyboardMarkup(keyboard=[[types.KeyboardButton("Mavzulashtirilgan")], [types.KeyboardButton("Attestatsiya")], [types.KeyboardButton("🔙 Ortga")]], resize_keyboard=True)
    await message.answer("Test turini tanlang!:", reply_markup=keyboard)
    await state.set_state("admin:create_test_type")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:create_test_type", user_id=ADMINS)
async def admin_create_test_time(message: types.Message, state: FSMContext):
    test_type = message.text
    keyboard = types.ReplyKeyboardMarkup(keyboard=[[types.KeyboardButton("Mavzulashtirilgan")], [types.KeyboardButton("Attestatsiya")], [types.KeyboardButton("🔙 Ortga")]], resize_keyboard=True)

    if test_type not in ["Mavzulashtirilgan", "Attestatsiya"]:
        await message.answer("Test turini tanlang!:", reply_markup=keyboard)
        return
    elif test_type == "Mavzulashtirilgan":
        test_type = "themed"
    elif test_type == "Attestatsiya":
        test_type = "attestation"
    data = await state.get_data()
    subject = data["subject"]
    test_name = data["test_name"]
    shuffling_type = data["shuffling_type"]
    timer = data['timer']
    test_id = testsdb.create_test(subject, test_name, shuffling_type, timer, test_type=test_type)
    await state.set_data({"subject": subject, "test_name": test_id})
    await message.answer("Test turi tanlandi!", reply_markup=ReplyKeyboardRemove())
    await message.answer("Test yaratildi! Testni savollarini belgilangan formatda menga yuklang!(WORD shaklida)", reply_markup=admin_back_book_list(subject, "tests"))
    await state.set_state("admin:create_test_upload")



            

# admin:create_test_upload
@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:create_test_upload", user_id=ADMINS)
async def admin_create_test_upload(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data.get("subject", "Aniqmas")
    test_name = data.get("test_name", "Aniqmas")
    document = message.document
    if not document.file_name.endswith(".docx"):
        await message.answer("Iltimos, faqat .docx formatdagi faylni yuboring!")
        await state.set_state("admin:create_test_upload")
        return
    
    file_bytes = BytesIO()
    await document.download(destination=file_bytes)
    file_bytes.seek(0)  

    doc = Document(file_bytes)
    text = "\n".join([para.text for para in doc.paragraphs])
    testsdb.parse_and_save_test(test_name, text)

    test_details = testsdb.get_information(test_name)
    await message.answer(f"<b>✍️ Test nomi:</b> {test_details['test_name']}\n<b>🔀 Test aralashtirish turi:</b> {test_details['shuffling_type']}\n<b>📍 Test savollar soni:</b> {test_details['total_questions']}\n<b>📅 Test yaratilgan sana va vaqt:</b> {test_details['created_time']}\n<b>⏰ Test vaqti:</b> {test_details['timer']}\n\n<b>🔗 Testni tarqatish uchun link:</b> https://t.me/PedagogUzbot?start=test{test_name}", reply_markup=test_details_keyboard(subject, test_name))





@dp.callback_query_handler(IsPrivate(), text_contains="admin:change_shuffling_type", state="*", user_id=ADMINS)
async def admin_change_shuffling_type(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    test_name = callback.data.split(":")[3]
    await callback.message.delete()
    await callback.message.answer("Testni aralashtirish turini tanlang:", reply_markup=shuffling_type_keyboard)
    await state.set_data({"subject": subject, "test_name": test_name})
    await state.set_state("admin:change_shuffling_type")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:change_shuffling_type", user_id=ADMINS)
async def admin_change_shuffling_type_message(message: types.Message, state: FSMContext):
    shuffling_type = message.text
    data = await state.get_data()
    subject = data["subject"]
    test_name = data["test_name"]
    if shuffling_type not in ["🔀 Hammasini aralashtirish", "🔀 Savolni aralashtirish", "🔀 Javobni aralashtirish", "🔀 Umuman aralashtirmaslik"]:
        await message.answer("Testni aralashtirish turini tanlang!", reply_markup=shuffling_type_keyboard)
        return
    elif shuffling_type == "🔀 Hammasini aralashtirish":
        shuffling_type = "shuffle_all"
    elif shuffling_type == "🔀 Savolni aralashtirish":
        shuffling_type = "shuffle_question"
    elif shuffling_type == "🔀 Javobni aralashtirish":
        shuffling_type = "shuffle_answer"
    elif shuffling_type == "🔀 Umuman aralashtirmaslik":
        shuffling_type = "no_shuffle"
    testsdb.change_shuffling_type(test_name, shuffling_type)
    test_details = testsdb.get_information(test_name)
    await message.answer("Testni aralashtirish turi o'zgartirildi!", reply_markup=ReplyKeyboardRemove())
    await message.answer(f"<b>✍️ Test nomi:</b> {test_details['test_name']}\n<b>🔀 Test aralashtirish turi:</b> {test_details['shuffling_type']}\n<b>📍 Test savollar soni:</b> {test_details['total_questions']}\n<b>📅 Test yaratilgan sana va vaqt:</b> {test_details['created_time']}\n<b>⏰ Test vaqti:</b> {test_details['timer']}", reply_markup=test_details_keyboard(subject, test_name))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:change_time", state="*", user_id=ADMINS)
async def admin_change_time(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    test_name = callback.data.split(":")[3]
    await callback.message.delete()
    await callback.message.answer("Testni vaqtni kiriting (sekundda):", reply_markup=ReplyKeyboardRemove())
    await state.set_data({"subject": subject, "test_name": test_name})
    await state.set_state("admin:change_time")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.TEXT, state="admin:change_time", user_id=ADMINS)
async def admin_change_time_message(message: types.Message, state: FSMContext):
    time = message.text
    data = await state.get_data()
    subject = data["subject"]
    test_name = data["test_name"]
    testsdb.change_time(test_name, time)
    test_details = testsdb.get_information(test_name)
    await message.answer("Testni vaqti o'zgartirildi!", reply_markup=ReplyKeyboardRemove())
    await message.answer(f"<b>✍️ Test nomi:</b> {test_details['test_name']}\n<b>🔀 Test aralashtirish turi:</b> {test_details['shuffling_type']}\n<b>📍 Test savollar soni:</b> {test_details['total_questions']}\n<b>📅 Test yaratilgan sana va vaqt:</b> {test_details['created_time']}\n<b>⏰ Test vaqti:</b> {test_details['timer']}", reply_markup=test_details_keyboard(subject, test_name))

@dp.callback_query_handler(IsPrivate(), text_contains="admin:delete_test", state="*", user_id=ADMINS)
async def admin_delete_test(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    test_name = callback.data.split(":")[3]
    testsdb.delete_test(test_name)
    await callback.answer("Test Bazadan o'chirib tashlandi")
    await admin_tests_action(callback, state)


@dp.callback_query_handler(IsPrivate(), text_contains="admin:add_question", state="*", user_id=ADMINS)
async def admin_add_test(callback: types.CallbackQuery, state: FSMContext):
    await state.finish()
    subject = callback.data.split(":")[2]
    test_name = callback.data.split(":")[3]
    await callback.message.delete()
    await callback.message.answer("Menga savollar mavjud word shaklida faylni yuboring:", reply_markup=ReplyKeyboardRemove())
    await state.set_data({"subject": subject, "test_name": test_name})
    await state.set_state("admin:add_question")


@dp.callback_query_handler(state="*", text='admin:statistics')
async def bot_stat(call: types.CallbackQuery, state: FSMContext):
    await call.answer("Kutib turing ...")
    count = db.count_users()[0]
    await call.message.edit_text(f"🤖 Botda jami <b>{count}</b> ta obunachi, aktiv obunachilar malumoti tez orada keladi.", reply_markup=admin_back_keyboard)
    await state.finish()
    users = db.select_all_users()
    x = 0
    y = 0
    for user in users:
        try:
            await bot.get_chat(user[0])
            x += 1
        except:
            y += 1
    await call.message.answer(f"✅ Aktiv: {x}\n"
                      f"❌ Bloklangan: {y}\n"
                      f"➖➖➖➖➖➖\n"
                      f"Umumiy: {count} ta")

@dp.message_handler(IsPrivate(), content_types=types.ContentType.DOCUMENT, state="admin:add_question", user_id=ADMINS)
async def admin_add_question_document(message: types.Message, state: FSMContext):
    data = await state.get_data()
    await state.finish()
    subject = data["subject"]
    test_name = data["test_name"]
    document = message.document
    if not document.file_name.endswith(".docx"):
        await message.answer("Iltimos, faqat .docx formatdagi faylni yuboring!")
        await state.update_data({"subject": subject, "test_name": test_name})
        await state.set_state("admin:add_question")
        return
    
    file_bytes = BytesIO()
    await document.download(destination=file_bytes)
    file_bytes.seek(0)  

    doc = Document(file_bytes)
    text = "\n".join([para.text for para in doc.paragraphs])
    testsdb.parse_and_save_test(test_name, text)
    await message.answer("Savollar yuklandi!", reply_markup=test_details_keyboard(subject, test_name))

# @dp.callback_query_handler(IsPrivate(), text_contains="admin:back_tests_page", state="*", user_id=ADMINS)
